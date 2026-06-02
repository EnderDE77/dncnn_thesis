import json
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.gridspec as gridspec
import os

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("python-docx not installed — tables saved as CSV only")

try:
    import csv
    CSV_AVAILABLE = True
except ImportError:
    CSV_AVAILABLE = False

PUBLISHED = {
    'BM3D':   {15: 31.73, 25: 28.61, 50: 25.62},
    'DnCNN':  {15: 31.73, 25: 29.23, 50: 26.23},
    'FFDNet': {15: 31.63, 25: 29.19, 50: 26.05},
}

METRICS = ['psnr', 'ssim', 'rmse', 'mae']
METRIC_LABELS = {
    'psnr': 'PSNR (dB)',
    'ssim': 'SSIM',
    'rmse': 'RMSE',
    'mae':  'MAE',
    'train_loss': 'Training Loss (MSE)'
}

def load_all_results():
    path = 'results/all_results.json'
    if not os.path.exists(path):
        print(f"Results not found at {path}")
        return {}
    with open(path) as f:
        return json.load(f)

def filter_results(results, sigma=None, domain=None,
                   kernel=None, activation=None, batch_size=None):
    filtered = {}
    for name, log in results.items():
        if sigma and log.get('sigma') != sigma:
            continue
        if domain and log.get('domain') != domain:
            continue
        if kernel and log.get('kernel_size') != kernel:
            continue
        if activation and log.get('activation') != activation:
            continue
        if batch_size and log.get('batch_size') != batch_size:
            continue
        filtered[name] = log
    return filtered

def short_name(config_name):
    parts = config_name.split('_')
    short = []
    for p in parts:
        if p.startswith('sigma'):
            short.append(p)
        elif p.startswith('k'):
            short.append(p)
        elif p in ['relu', 'leakyrelu']:
            short.append('LReLU' if p == 'leakyrelu' else 'ReLU')
        elif p.startswith('bs'):
            short.append(p)
        elif p in ['image', 'fourier']:
            short.append('Img' if p == 'image' else 'FFT')
        elif p == 'complex':
            pass
    return ' '.join(short)

# ─── 1. Training curves per metric ───────────────────────────────────────────

def plot_training_curves_by_sigma(results):
    for sigma in [15, 25, 50]:
        subset = filter_results(results, sigma=sigma)
        if not subset:
            continue

        fig, axes = plt.subplots(2, 3, figsize=(18, 10))
        axes = axes.flatten()

        plot_metrics = ['train_loss', 'val_psnr',
                        'val_ssim', 'val_rmse',
                        'val_mae', 'set12_psnr']
        plot_labels  = ['Train Loss', 'Val PSNR',
                        'Val SSIM',  'Val RMSE',
                        'Val MAE',   'Set12 PSNR']

        colors = plt.cm.tab20(np.linspace(0, 1, len(subset)))

        for idx, (metric, label) in enumerate(
                zip(plot_metrics, plot_labels)):
            ax = axes[idx]
            for (name, log), color in zip(subset.items(), colors):
                if metric in log:
                    ax.plot(log['epochs'], log[metric],
                            color=color,
                            label=short_name(name),
                            linewidth=1.2)
            ax.set_xlabel('Epoch')
            ax.set_ylabel(label)
            ax.set_title(label)
            ax.grid(True, alpha=0.3)
            if idx == 0:
                ax.legend(fontsize=6, ncol=2,
                          loc='upper right')

        fig.suptitle(f'Training Curves — σ={sigma} '
                     f'(All Combinations)', fontsize=14)
        plt.tight_layout()
        path = f'results/plots/curves_sigma{sigma}.png'
        plt.savefig(path, dpi=150, bbox_inches='tight')
        plt.close()
        print(f"Saved: {path}")

def plot_training_curves_by_domain(results):
    for domain in ['image', 'fourier_complex']:
        subset = filter_results(results, domain=domain)
        if not subset:
            continue

        fig, axes = plt.subplots(2, 3, figsize=(18, 10))
        axes = axes.flatten()

        plot_metrics = ['train_loss', 'val_psnr',
                        'val_ssim', 'val_rmse',
                        'val_mae', 'set12_psnr']
        plot_labels  = ['Train Loss', 'Val PSNR',
                        'Val SSIM',  'Val RMSE',
                        'Val MAE',   'Set12 PSNR']

        colors = plt.cm.tab20(np.linspace(0, 1, len(subset)))

        for idx, (metric, label) in enumerate(
                zip(plot_metrics, plot_labels)):
            ax = axes[idx]
            for (name, log), color in zip(subset.items(), colors):
                if metric in log:
                    ax.plot(log['epochs'], log[metric],
                            color=color,
                            label=short_name(name),
                            linewidth=1.2)
            ax.set_xlabel('Epoch')
            ax.set_ylabel(label)
            ax.set_title(label)
            ax.grid(True, alpha=0.3)

        domain_label = ('Image Domain'
                        if domain == 'image'
                        else 'Fourier Complex Domain')
        fig.suptitle(f'Training Curves — {domain_label} '
                     f'(All Combinations)', fontsize=14)
        plt.tight_layout()
        path = (f'results/plots/'
                f'curves_{domain}.png')
        plt.savefig(path, dpi=150, bbox_inches='tight')
        plt.close()
        print(f"Saved: {path}")

# ─── 2. Bar charts ────────────────────────────────────────────────────────────

def plot_best_per_combination(results):
    """Bar chart of best PSNR per sigma per domain."""
    for sigma in [15, 25, 50]:
        for domain in ['image', 'fourier_complex']:
            subset = filter_results(
                results, sigma=sigma, domain=domain)
            if not subset:
                continue

            names, psnrs, ssims, rmses, maes, runtimes = \
                [], [], [], [], [], []
            for name, log in subset.items():
                fs = log.get('final_set12', {})
                names.append(short_name(name))
                psnrs.append(fs.get('psnr', 0))
                ssims.append(fs.get('ssim', 0))
                rmses.append(fs.get('rmse', 0))
                maes.append(fs.get('mae', 0))
                runtimes.append(
                    log.get('total_runtime_minutes', 0))

            x = np.arange(len(names))
            fig, axes = plt.subplots(2, 2, figsize=(16, 10))

            data_sets = [
                (psnrs,    'PSNR (dB)',          axes[0, 0]),
                (ssims,    'SSIM',                axes[0, 1]),
                (rmses,    'RMSE (lower=better)', axes[1, 0]),
                (runtimes, 'Runtime (minutes)',   axes[1, 1]),
            ]

            for vals, ylabel, ax in data_sets:
                bars = ax.bar(x, vals,
                              color=plt.cm.tab20(
                                  np.linspace(0, 1, len(names))))
                ax.set_xticks(x)
                ax.set_xticklabels(names,
                                   rotation=45, ha='right',
                                   fontsize=8)
                ax.set_ylabel(ylabel)
                ax.set_title(ylabel)
                ax.grid(True, axis='y', alpha=0.3)
                for bar in bars:
                    h = bar.get_height()
                    ax.text(bar.get_x() + bar.get_width()/2,
                            h * 1.01,
                            f'{h:.3f}',
                            ha='center', va='bottom',
                            fontsize=7)

            domain_label = ('Image' if domain == 'image'
                            else 'Fourier Complex')
            fig.suptitle(
                f'All Metrics — σ={sigma}, '
                f'{domain_label} Domain',
                fontsize=13)
            plt.tight_layout()
            path = (f'results/plots/'
                    f'bars_sigma{sigma}_{domain}.png')
            plt.savefig(path, dpi=150, bbox_inches='tight')
            plt.close()
            print(f"Saved: {path}")

def plot_domain_comparison(results):
    """Image vs Fourier for best config per sigma."""
    for sigma in [15, 25, 50]:
        img_sub = filter_results(
            results, sigma=sigma, domain='image')
        fft_sub = filter_results(
            results, sigma=sigma, domain='fourier_complex')
        if not img_sub or not fft_sub:
            continue

        # Best of each domain by Set12 PSNR
        best_img = max(
            img_sub.items(),
            key=lambda x: x[1].get(
                'final_set12', {}).get('psnr', 0))
        best_fft = max(
            fft_sub.items(),
            key=lambda x: x[1].get(
                'final_set12', {}).get('psnr', 0))

        metrics = ['psnr', 'ssim', 'rmse', 'mae']
        labels  = ['PSNR (dB)', 'SSIM', 'RMSE', 'MAE']
        img_vals = [best_img[1]['final_set12'].get(m, 0)
                    for m in metrics]
        fft_vals = [best_fft[1]['final_set12'].get(m, 0)
                    for m in metrics]

        x = np.arange(len(metrics))
        fig, ax = plt.subplots(figsize=(10, 6))
        w = 0.35
        bars1 = ax.bar(x - w/2, img_vals, w,
                       label='Image Domain', color='steelblue')
        bars2 = ax.bar(x + w/2, fft_vals, w,
                       label='Fourier Complex', color='darkorange')
        ax.set_xticks(x)
        ax.set_xticklabels(labels)
        ax.set_title(f'Best Image vs Fourier — σ={sigma} '
                     f'(Set12)')
        ax.legend()
        ax.grid(True, axis='y', alpha=0.3)
        for bar in bars1 + bars2:
            h = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2,
                    h * 1.01, f'{h:.3f}',
                    ha='center', va='bottom', fontsize=8)
        plt.tight_layout()
        path = f'results/plots/domain_compare_sigma{sigma}.png'
        plt.savefig(path, dpi=150, bbox_inches='tight')
        plt.close()
        print(f"Saved: {path}")

def plot_literature_comparison(results):
    """Compare best results vs published benchmarks."""
    fig, axes = plt.subplots(1, 3, figsize=(15, 6))

    for idx, sigma in enumerate([15, 25, 50]):
        ax = axes[idx]
        methods, psnrs = [], []

        # Published baselines
        for method, vals in PUBLISHED.items():
            methods.append(method)
            psnrs.append(vals[sigma])

        # Our best image domain result
        img_sub = filter_results(
            results, sigma=sigma, domain='image')
        if img_sub:
            best = max(
                img_sub.items(),
                key=lambda x: x[1].get(
                    'final_val', {}).get('psnr', 0))
            methods.append('Ours (Image)')
            psnrs.append(
                best[1].get('final_val', {}).get('psnr', 0))

        # Our best Fourier result
        fft_sub = filter_results(
            results, sigma=sigma, domain='fourier_complex')
        if fft_sub:
            best = max(
                fft_sub.items(),
                key=lambda x: x[1].get(
                    'final_val', {}).get('psnr', 0))
            methods.append('Ours (Fourier)')
            psnrs.append(
                best[1].get('final_val', {}).get('psnr', 0))

        colors = (['gray'] * len(PUBLISHED) +
                  ['steelblue', 'darkorange'])
        bars = ax.bar(methods, psnrs, color=colors[:len(methods)])
        ax.set_title(f'σ={sigma}')
        ax.set_ylabel('BSD68 PSNR (dB)')
        ax.set_xticklabels(methods, rotation=20, ha='right')
        ax.grid(True, axis='y', alpha=0.3)
        for bar in bars:
            h = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2,
                    h + 0.05, f'{h:.2f}',
                    ha='center', va='bottom', fontsize=8)

    fig.suptitle('Comparison with Published Benchmarks — '
                 'BSD68 PSNR (dB)', fontsize=13)
    plt.tight_layout()
    plt.savefig('results/plots/literature_comparison.png',
                dpi=150, bbox_inches='tight')
    plt.close()
    print("Saved: results/plots/literature_comparison.png")

def plot_batch_size_effect(results):
    """Show effect of batch size on PSNR per domain."""
    for sigma in [15, 25, 50]:
        for domain in ['image', 'fourier_complex']:
            fig, ax = plt.subplots(figsize=(10, 6))
            plotted = False

            for kernel in [3, 5]:
                for activation in ['relu', 'leakyrelu']:
                    bs_list, psnr_list = [], []
                    for bs in [32, 64]:
                        subset = filter_results(
                            results,
                            sigma=sigma,
                            domain=domain,
                            kernel=kernel,
                            activation=activation,
                            batch_size=bs)
                        if subset:
                            log = list(subset.values())[0]
                            bs_list.append(bs)
                            psnr_list.append(
                                log.get('final_set12',
                                        {}).get('psnr', 0))
                    if bs_list:
                        label = (f'k{kernel}x{kernel} '
                                 f'{"LReLU" if activation == "leakyrelu" else "ReLU"}')
                        ax.plot(bs_list, psnr_list,
                                marker='o', label=label)
                        plotted = True

            if plotted:
                ax.set_xlabel('Batch Size')
                ax.set_ylabel('Set12 PSNR (dB)')
                domain_label = ('Image'
                                if domain == 'image'
                                else 'Fourier Complex')
                ax.set_title(
                    f'Batch Size Effect — σ={sigma}, '
                    f'{domain_label}')
                ax.legend()
                ax.grid(True, alpha=0.3)
                plt.tight_layout()
                path = (f'results/plots/'
                        f'batchsize_sigma{sigma}_{domain}.png')
                plt.savefig(path, dpi=150, bbox_inches='tight')
                plt.close()
                print(f"Saved: {path}")

# ─── 3. Tables ────────────────────────────────────────────────────────────────

def generate_summary_csv(results):
    import csv
    path = 'results/tables/summary_all_results.csv'
    os.makedirs('results/tables', exist_ok=True)

    with open(path, 'w', newline='') as f:
        writer = csv.writer(f)
        writer.writerow([
            'Name', 'Sigma', 'Kernel', 'Activation',
            'Batch Size', 'Domain',
            'Val PSNR', 'Val SSIM', 'Val RMSE', 'Val MAE',
            'Set12 PSNR', 'Set12 SSIM', 'Set12 RMSE', 'Set12 MAE',
            'Runtime (min)'
        ])
        for name, log in results.items():
            fv = log.get('final_val', {})
            fs = log.get('final_set12', {})
            writer.writerow([
                name,
                log.get('sigma', ''),
                log.get('kernel_size', ''),
                log.get('activation', ''),
                log.get('batch_size', ''),
                log.get('domain', ''),
                f"{fv.get('psnr', 0):.4f}",
                f"{fv.get('ssim', 0):.4f}",
                f"{fv.get('rmse', 0):.5f}",
                f"{fv.get('mae', 0):.5f}",
                f"{fs.get('psnr', 0):.4f}",
                f"{fs.get('ssim', 0):.4f}",
                f"{fs.get('rmse', 0):.5f}",
                f"{fs.get('mae', 0):.5f}",
                f"{log.get('total_runtime_minutes', 0):.1f}",
            ])
    print(f"Saved: {path}")

def generate_best_per_sigma_csv(results):
    import csv
    path = 'results/tables/best_per_sigma.csv'
    os.makedirs('results/tables', exist_ok=True)

    with open(path, 'w', newline='') as f:
        writer = csv.writer(f)
        writer.writerow([
            'Sigma', 'Domain', 'Best Config',
            'Val PSNR', 'Set12 PSNR', 'Set12 SSIM',
            'Set12 RMSE', 'Set12 MAE', 'Runtime (min)'
        ])
        for sigma in [15, 25, 50]:
            for domain in ['image', 'fourier_complex']:
                subset = filter_results(
                    results, sigma=sigma, domain=domain)
                if not subset:
                    continue
                best_name, best_log = max(
                    subset.items(),
                    key=lambda x: x[1].get(
                        'final_set12', {}).get('psnr', 0))
                fs = best_log.get('final_set12', {})
                writer.writerow([
                    sigma,
                    domain,
                    short_name(best_name),
                    f"{best_log.get('final_val', {}).get('psnr', 0):.4f}",
                    f"{fs.get('psnr', 0):.4f}",
                    f"{fs.get('ssim', 0):.4f}",
                    f"{fs.get('rmse', 0):.5f}",
                    f"{fs.get('mae', 0):.5f}",
                    f"{best_log.get('total_runtime_minutes', 0):.1f}",
                ])
    print(f"Saved: {path}")

def generate_literature_comparison_csv(results):
    import csv
    path = 'results/tables/literature_comparison.csv'
    os.makedirs('results/tables', exist_ok=True)

    with open(path, 'w', newline='') as f:
        writer = csv.writer(f)
        writer.writerow([
            'Method', 'Sigma=15 PSNR',
            'Sigma=25 PSNR', 'Sigma=50 PSNR', 'Sigma=100 PSNR', 'Source'
        ])
        for method, vals in PUBLISHED.items():
            writer.writerow([
                method,
                f"{vals[15]:.2f}",
                f"{vals[25]:.2f}",
                f"{vals[50]:.2f}",
                f"{vals[100]:.2f}",
                'Literature'
            ])
        for domain in ['image', 'fourier_complex']:
            row = [
                f"Ours ({'Image' if domain == 'image' else 'Fourier'})"]
            for sigma in [15, 25, 50, 100]:
                subset = filter_results(
                    results, sigma=sigma, domain=domain)
                if subset:
                    best = max(
                        subset.items(),
                        key=lambda x: x[1].get(
                            'final_val', {}).get('psnr', 0))
                    row.append(
                        f"{best[1].get('final_val', {}).get('psnr', 0):.2f}")
                else:
                    row.append('N/A')
            row.append('This work')
            writer.writerow(row)
    print(f"Saved: {path}")

def generate_word_tables(results):
    if not DOCX_AVAILABLE:
        print("Skipping Word tables — python-docx not installed")
        return

    doc = Document()
    doc.add_heading('DnCNN Experiment Results', level=1)

    # Table 1 — Literature comparison
    doc.add_heading('Table 1: Comparison with Published Benchmarks',
                    level=2)
    t1 = doc.add_table(rows=1, cols=5)
    t1.style = 'Table Grid'
    h = t1.rows[0].cells
    for i, text in enumerate([
            'Method', 'σ=15 PSNR', 'σ=25 PSNR',
            'σ=50 PSNR', 'Source']):
        h[i].text = text
        h[i].paragraphs[0].runs[0].bold = True

    for method, vals in PUBLISHED.items():
        r = t1.add_row().cells
        r[0].text = method
        r[1].text = f"{vals[15]:.2f}"
        r[2].text = f"{vals[25]:.2f}"
        r[3].text = f"{vals[50]:.2f}"
        r[3].text = f"{vals[100]:.2f}"
        r[5].text = 'Literature'

    for domain in ['image', 'fourier_complex']:
        r = t1.add_row().cells
        r[0].text = (f"Ours "
                     f"({'Image' if domain == 'image' else 'Fourier'})")
        for col_idx, sigma in enumerate([15, 25, 50], start=1):
            subset = filter_results(
                results, sigma=sigma, domain=domain)
            if subset:
                best = max(
                    subset.items(),
                    key=lambda x: x[1].get(
                        'final_val', {}).get('psnr', 0))
                r[col_idx].text = (
                    f"{best[1].get('final_val', {}).get('psnr', 0):.2f}")
            else:
                r[col_idx].text = 'N/A'
        r[4].text = 'This work'

    doc.add_paragraph()

    # Table 2 — Best per sigma
    doc.add_heading('Table 2: Best Configuration per Sigma and Domain',
                    level=2)
    t2 = doc.add_table(rows=1, cols=9)
    t2.style = 'Table Grid'
    h2 = t2.rows[0].cells
    for i, text in enumerate([
            'Sigma', 'Domain', 'Best Config',
            'Val PSNR', 'Set12 PSNR', 'Set12 SSIM',
            'RMSE', 'MAE', 'Runtime (min)']):
        h2[i].text = text
        h2[i].paragraphs[0].runs[0].bold = True

    for sigma in [15, 25, 50]:
        for domain in ['image', 'fourier_complex']:
            subset = filter_results(
                results, sigma=sigma, domain=domain)
            if not subset:
                continue
            best_name, best_log = max(
                subset.items(),
                key=lambda x: x[1].get(
                    'final_set12', {}).get('psnr', 0))
            fs = best_log.get('final_set12', {})
            fv = best_log.get('final_val', {})
            r = t2.add_row().cells
            r[0].text = str(sigma)
            r[1].text = ('Image'
                         if domain == 'image' else 'Fourier')
            r[2].text = short_name(best_name)
            r[3].text = f"{fv.get('psnr', 0):.2f}"
            r[4].text = f"{fs.get('psnr', 0):.2f}"
            r[5].text = f"{fs.get('ssim', 0):.4f}"
            r[6].text = f"{fs.get('rmse', 0):.5f}"
            r[7].text = f"{fs.get('mae', 0):.5f}"
            r[8].text = (
                f"{best_log.get('total_runtime_minutes', 0):.1f}")

    doc.add_paragraph()

    # Table 3 — Full results
    doc.add_heading('Table 3: Full Results — All Configurations',
                    level=2)
    t3 = doc.add_table(rows=1, cols=10)
    t3.style = 'Table Grid'
    h3 = t3.rows[0].cells
    for i, text in enumerate([
            'Config', 'σ', 'Kernel', 'Act', 'BS', 'Domain',
            'PSNR', 'SSIM', 'RMSE', 'Runtime']):
        h3[i].text = text
        h3[i].paragraphs[0].runs[0].bold = True

    for name, log in results.items():
        fs = log.get('final_set12', {})
        r = t3.add_row().cells
        r[0].text = short_name(name)
        r[1].text = str(log.get('sigma', ''))
        r[2].text = (f"{log.get('kernel_size', '')}x"
                     f"{log.get('kernel_size', '')}")
        r[3].text = ('LReLU'
                     if log.get('activation') == 'leakyrelu'
                     else 'ReLU')
        r[4].text = str(log.get('batch_size', ''))
        r[5].text = ('Image'
                     if log.get('domain') == 'image'
                     else 'Fourier')
        r[6].text = f"{fs.get('psnr', 0):.2f}"
        r[7].text = f"{fs.get('ssim', 0):.4f}"
        r[8].text = f"{fs.get('rmse', 0):.5f}"
        r[9].text = (
            f"{log.get('total_runtime_minutes', 0):.1f}m")

    os.makedirs('results/tables', exist_ok=True)
    doc.save('results/tables/all_experiment_tables.docx')
    print("Saved: results/tables/all_experiment_tables.docx")

# ─── 4. Summary print ─────────────────────────────────────────────────────────

def print_summary(results):
    print("\n" + "="*80)
    print("EXPERIMENT SUMMARY")
    print("="*80)
    print(f"{'Config':<35} {'σ':>4} {'Domain':<8} "
          f"{'Val PSNR':>9} {'S12 PSNR':>9} "
          f"{'SSIM':>7} {'Runtime':>9}")
    print("-"*80)
    for name, log in sorted(
            results.items(),
            key=lambda x: x[1].get(
                'final_set12', {}).get('psnr', 0),
            reverse=True):
        fs = log.get('final_set12', {})
        fv = log.get('final_val', {})
        domain_s = ('Image'
                    if log.get('domain') == 'image'
                    else 'Fourier')
        print(f"{short_name(name):<35} "
              f"{log.get('sigma', ''):>4} "
              f"{domain_s:<8} "
              f"{fv.get('psnr', 0):>9.2f} "
              f"{fs.get('psnr', 0):>9.2f} "
              f"{fs.get('ssim', 0):>7.4f} "
              f"{log.get('total_runtime_minutes', 0):>8.1f}m")
    print("="*80)

# ─── Main ─────────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    os.makedirs('results/plots', exist_ok=True)
    os.makedirs('results/tables', exist_ok=True)

    results = load_all_results()
    if not results:
        print("No results found. Run train_all.py first.")
    else:
        print(f"Loaded {len(results)} experiment results\n")

        print("Generating training curves by sigma...")
        plot_training_curves_by_sigma(results)

        print("Generating training curves by domain...")
        plot_training_curves_by_domain(results)

        print("Generating best per combination bar charts...")
        plot_best_per_combination(results)

        print("Generating domain comparison charts...")
        plot_domain_comparison(results)

        print("Generating literature comparison chart...")
        plot_literature_comparison(results)

        print("Generating batch size effect charts...")
        plot_batch_size_effect(results)

        print("Generating CSV tables...")
        generate_summary_csv(results)
        generate_best_per_sigma_csv(results)
        generate_literature_comparison_csv(results)

        print("Generating Word tables...")
        generate_word_tables(results)

        print_summary(results)

        print("\nAll plots and tables generated.")
        print("Plots: results/plots/")
        print("Tables: results/tables/")